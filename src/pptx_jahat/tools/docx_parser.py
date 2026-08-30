import docx
from pathlib import Path
from typing import Dict, Any, List

def parse_docx(file_path: Path | str) -> Dict[str, Any]:
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Docx file not found: {file_path}")
        
    doc = docx.Document(str(path))
    
    raw_paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    sections = []
    current_section = {
        "title": "مقدمه",
        "level": 1,
        "paragraphs": [],
        "bullets": [],
        "tables": []
    }
    
    doc_title = path.stem.replace("_", " ")
    has_found_heading = False
    
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
            
        style_name = p.style.name.lower() if p.style else ""
        
        # Check if heading or title or bold/special
        if "heading 1" in style_name or "title" in style_name or text.startswith("فصل") or text.startswith("درس"):
            if not has_found_heading:
                doc_title = text
                has_found_heading = True
            if current_section["paragraphs"] or current_section["bullets"] or current_section["tables"]:
                sections.append(current_section)
            current_section = {
                "title": text,
                "level": 1,
                "paragraphs": [],
                "bullets": [],
                "tables": []
            }
        elif "heading 2" in style_name or "heading 3" in style_name or text.startswith("راهبرد") or text.startswith("سوال") or text.startswith("مسئله") or text.startswith("تمرین") or text.startswith("نکته"):
            if current_section["paragraphs"] or current_section["bullets"] or current_section["tables"]:
                sections.append(current_section)
            current_section = {
                "title": text,
                "level": 2 if "heading 2" in style_name else 3,
                "paragraphs": [],
                "bullets": [],
                "tables": []
            }
        elif "list" in style_name or "bullet" in style_name or text.startswith("- ") or text.startswith("• ") or text.startswith("* "):
            cleaned = text.lstrip("- •*").strip()
            current_section["bullets"].append(cleaned)
        else:
            current_section["paragraphs"].append(text)
            
    # Tables parsing
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_vals = [cell.text.strip() for cell in row.cells]
            table_data.append(row_vals)
        if table_data:
            current_section["tables"].append(table_data)
            
    if current_section["paragraphs"] or current_section["bullets"] or current_section["tables"]:
        sections.append(current_section)
        
    if not sections:
        sections.append({
            "title": doc_title,
            "level": 1,
            "paragraphs": raw_paragraphs or ["Content imported from document."],
            "bullets": [],
            "tables": []
        })
        
    return {
        "document_title": doc_title,
        "total_sections": len(sections),
        "raw_paragraphs": raw_paragraphs,
        "sections": sections
    }
